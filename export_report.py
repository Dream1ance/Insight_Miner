
import os
import sys
import argparse

# Try to import pdfkit first (preferred - preserves exact HTML styling)
try:
    import pdfkit
    PDFKIT_AVAILABLE = True
except ImportError:
    PDFKIT_AVAILABLE = False

# Fallback to docx2pdf if pdfkit not available
try:
    from docx2pdf import convert as _docx2pdf_convert
except Exception:
    _docx2pdf_convert = None

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt


def html_to_word(html_file_or_string, docx_file, is_path=True):
    """Convert HTML (file path or raw HTML string) to a .docx file.

    Args:
        html_file_or_string: path to an HTML file or a raw HTML string.
        docx_file: output .docx path.
        is_path: if True treat first arg as a filepath, otherwise as HTML content.
    """
    if is_path:
        with open(html_file_or_string, "r", encoding="utf-8") as f:
            html_content = f.read()
    else:
        html_content = html_file_or_string

    soup = BeautifulSoup(html_content, "html.parser")

    doc = Document()

    # Optional: add a title if present
    title_tag = soup.find(['h1', 'title'])
    if title_tag and title_tag.get_text(strip=True):
        doc.add_heading(title_tag.get_text(strip=True), level=1)

    # Very small formatting handling: paragraphs and headings
    for elem in soup.find_all(['h1', 'h2', 'h3', 'p', 'li']):
        text = elem.get_text(strip=True)
        if not text:
            continue
        if elem.name == 'h1':
            doc.add_heading(text, level=1)
        elif elem.name == 'h2':
            doc.add_heading(text, level=2)
        elif elem.name == 'h3':
            doc.add_heading(text, level=3)
        elif elem.name == 'li':
            # add as a simple paragraph prefixed with a bullet
            p = doc.add_paragraph('\u2022 ' + text)
            # set font size a bit smaller
            for run in p.runs:
                run.font.size = Pt(11)
        else:
            doc.add_paragraph(text)

    # Save
    doc.save(docx_file)
    print("DOCX created:", docx_file)


def word_to_pdf(docx_file, pdf_file):
    """Convert a .docx file into a .pdf using `docx2pdf`.

    On Windows, this uses Microsoft Word via COM and requires Word installed.
    """
    if _docx2pdf_convert is None:
        raise RuntimeError("docx2pdf is not installed or could not be imported. Install with `pip install docx2pdf`.")

    # docx2pdf.convert accepts either (input, output) or a single path
    _docx2pdf_convert(docx_file, pdf_file)
    print("PDF created:", pdf_file)


def html_to_pdf(html_file_or_string, pdf_file, is_path=True):
   
    if PDFKIT_AVAILABLE:
        # Use pdfkit - preserves exact HTML/CSS styling
        try:
            if is_path:
                # HTML file path
                pdfkit.from_file(html_file_or_string, pdf_file)
            else:
                # Raw HTML string
                pdfkit.from_string(html_file_or_string, pdf_file)
            print("PDF created (via pdfkit - exact HTML styling preserved):", pdf_file)
            return
        except Exception as e:
            print(f"pdfkit conversion failed: {e}. Falling back to DOCX method...")
    
    # Fallback: Convert HTML -> DOCX -> PDF
    if _docx2pdf_convert is None:
        raise RuntimeError(
            "PDF conversion failed. Install wkhtmltopdf for best results:\n"
            "  Windows: choco install wkhtmltopdf  OR  download from https://wkhtmltopdf.org\n"
            "  Linux: apt-get install wkhtmltopdf\n"
            "  macOS: brew install --cask wkhtmltopdf\n"
            "Then install pdfkit: pip install pdfkit"
        )
    
    # Fallback method: HTML -> DOCX -> PDF
    temp_docx = pdf_file.replace('.pdf', '_temp.docx')
    try:
        html_to_word(html_file_or_string, temp_docx, is_path=is_path)
        _docx2pdf_convert(temp_docx, pdf_file)
        print("PDF created (via DOCX fallback - styling may differ):", pdf_file)
    finally:
        if os.path.exists(temp_docx):
            try:
                os.remove(temp_docx)
            except Exception:
                pass


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Convert HTML to DOCX and PDF')
    parser.add_argument('--input', '-i', required=True, help='Input HTML file path')
    parser.add_argument('--docx', '-d', default='output.docx', help='Output DOCX path')
    parser.add_argument('--pdf', '-p', default='output.pdf', help='Output PDF path')

    args = parser.parse_args()

    # Run conversion (input is a path by default)
    try:
        html_to_word(args.input, args.docx, is_path=True)
        # Only attempt PDF conversion if docx2pdf is available
        if _docx2pdf_convert is None:
            print("Skipping PDF conversion: `docx2pdf` not available. Install with `pip install docx2pdf` and ensure MS Word is installed on Windows.")
            sys.exit(0)
        word_to_pdf(args.docx, args.pdf)
    except Exception as e:
        print("Conversion failed:", e)
        sys.exit(1)
